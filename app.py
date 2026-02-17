"""
Enhanced Voice Summarizer Backend
Flask API with Whisper, multi-language support, and advanced features
Compatible with Python 3.10, 3.11, 3.12, 3.13
"""

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from flask_socketio import SocketIO, emit
import os
import uuid
import json
from datetime import datetime
from pathlib import Path
import io
import wave
import re
from collections import Counter

# NumPy - safe import
try:
    import numpy as np
except ImportError:
    print("WARNING: numpy not installed. Run: pip install numpy")
    np = None

# Whisper - safe import
try:
    import whisper
    WHISPER_AVAILABLE = True
except ImportError:
    print("WARNING: whisper not installed. Run: pip install openai-whisper")
    WHISPER_AVAILABLE = False

# Librosa - safe import (replaces pydub for audio processing)
try:
    import librosa
    import soundfile as sf
    LIBROSA_AVAILABLE = True
except ImportError:
    print("WARNING: librosa not installed. Run: pip install librosa soundfile")
    LIBROSA_AVAILABLE = False

# Transformers - safe import
try:
    from transformers import pipeline
    TRANSFORMERS_AVAILABLE = True
except ImportError:
    print("WARNING: transformers not installed. Run: pip install transformers")
    TRANSFORMERS_AVAILABLE = False

# Language detection - safe import
try:
    from langdetect import detect, detect_langs
    LANGDETECT_AVAILABLE = True
except ImportError:
    print("WARNING: langdetect not installed. Run: pip install langdetect")
    LANGDETECT_AVAILABLE = False

# Translation - try deep-translator first (Python 3.13 compatible), fallback to googletrans
GOOGLETRANS_AVAILABLE = False
translator = None
DEEP_TRANSLATOR_AVAILABLE = False

try:
    from deep_translator import GoogleTranslator
    DEEP_TRANSLATOR_AVAILABLE = True
    GOOGLETRANS_AVAILABLE = True  # reuse same flag
    print("Translation ready (deep-translator)")
except ImportError:
    try:
        from googletrans import Translator
        translator = Translator()
        GOOGLETRANS_AVAILABLE = True
        print("Translation ready (googletrans)")
    except ImportError:
        print("WARNING: No translation library. Run: pip install deep-translator")

# NLTK - safe import
try:
    import nltk
    from nltk.tokenize import sent_tokenize
    from nltk.corpus import stopwords
    NLTK_AVAILABLE = True
except ImportError:
    print("WARNING: nltk not installed. Run: pip install nltk")
    NLTK_AVAILABLE = False

# python-docx - safe import
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    print("WARNING: python-docx not installed. Run: pip install python-docx")
    DOCX_AVAILABLE = False

# ReportLab - safe import
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    print("WARNING: reportlab not installed. Run: pip install reportlab")
    REPORTLAB_AVAILABLE = False

# Download required NLTK data
if NLTK_AVAILABLE:
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        nltk.download('punkt')
    try:
        nltk.data.find('corpora/stopwords')
    except LookupError:
        nltk.download('stopwords')

app = Flask(__name__, static_folder='.', static_url_path='')
CORS(app, resources={r"/*": {"origins": "*"}})
socketio = SocketIO(app, cors_allowed_origins="*", async_mode='threading')

# Configuration
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
HISTORY_FILE = 'history.json'
ALLOWED_EXTENSIONS = {'mp3', 'wav', 'ogg', 'flac', 'm4a', 'webm', 'mp4'}

# Create necessary directories
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Load models (lazy loading)
whisper_model = None
summarizer = None
sentiment_analyzer = None

def get_whisper_model():
    """Lazy load Whisper model"""
    global whisper_model
    if not WHISPER_AVAILABLE:
        raise ImportError("Whisper not installed. Run: pip install openai-whisper")
    if whisper_model is None:
        whisper_model = whisper.load_model("base")  # Can be changed to "small", "medium", "large"
    return whisper_model

def get_summarizer():
    """Lazy load summarization model"""
    global summarizer
    if not TRANSFORMERS_AVAILABLE:
        raise ImportError("Transformers not installed. Run: pip install transformers")
    if summarizer is None:
        summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
    return summarizer

def get_sentiment_analyzer():
    """Lazy load sentiment analysis model"""
    global sentiment_analyzer
    if not TRANSFORMERS_AVAILABLE:
        raise ImportError("Transformers not installed. Run: pip install transformers")
    if sentiment_analyzer is None:
        sentiment_analyzer = pipeline("sentiment-analysis")
    return sentiment_analyzer

def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_audio_features(audio_path):
    """Extract audio features for visualization using librosa only (no pydub)"""
    if not LIBROSA_AVAILABLE:
        return {
            'waveform': [],
            'duration': 0,
            'sample_rate': 0,
            'tempo': 0,
            'avg_spectral_centroid': 0
        }
    try:
        # Load audio file using librosa (handles MP3, WAV, OGG etc)
        y, sr = librosa.load(audio_path, sr=None)

        # Calculate waveform data (downsample for web display)
        waveform = librosa.resample(y, orig_sr=sr, target_sr=100)

        # Calculate duration
        duration = librosa.get_duration(y=y, sr=sr)

        # Extract additional features
        tempo, _ = librosa.beat.beat_track(y=y, sr=sr)
        spectral_centroids = librosa.feature.spectral_centroid(y=y, sr=sr)[0]

        # Fix NumPy deprecation: extract scalar properly from array
        tempo_value = float(tempo.item()) if hasattr(tempo, 'item') else float(tempo[0]) if hasattr(tempo, '__len__') else float(tempo)

        return {
            'waveform': waveform.tolist()[:1000],  # Limit points for performance
            'duration': float(duration),
            'sample_rate': int(sr),
            'tempo': tempo_value,
            'avg_spectral_centroid': float(np.mean(spectral_centroids)) if np is not None else 0
        }
    except Exception as e:
        print(f"Error extracting audio features: {e}")
        return {
            'waveform': [],
            'duration': 0,
            'sample_rate': 0,
            'tempo': 0,
            'avg_spectral_centroid': 0
        }

def detect_language_advanced(text):
    """Detect language with confidence scores"""
    if not LANGDETECT_AVAILABLE:
        return [{'language': 'unknown', 'confidence': 0}]
    try:
        langs = detect_langs(text)
        return [{'language': lang.lang, 'confidence': lang.prob} for lang in langs]
    except:
        return [{'language': 'unknown', 'confidence': 0}]

def extract_keywords(text, num_keywords=10):
    """Extract key terms from text"""
    try:
        # Remove punctuation and convert to lowercase
        text_clean = re.sub(r'[^\w\s]', '', text.lower())
        words = text_clean.split()

        # Remove stopwords if NLTK is available
        if NLTK_AVAILABLE:
            try:
                stop_words = set(stopwords.words('english'))
                words = [w for w in words if w not in stop_words and len(w) > 3]
            except:
                words = [w for w in words if len(w) > 3]
        else:
            # Basic stopwords fallback
            basic_stops = {'this', 'that', 'with', 'have', 'will', 'from', 'they',
                           'been', 'were', 'what', 'when', 'your', 'more', 'also'}
            words = [w for w in words if w not in basic_stops and len(w) > 3]

        word_freq = Counter(words)
        return [{'word': word, 'count': count} for word, count in word_freq.most_common(num_keywords)]
    except:
        return []

def generate_summary(text, summary_type='brief'):
    """Generate different types of summaries"""
    try:
        if not TRANSFORMERS_AVAILABLE:
            # Fallback: simple extractive summary
            sentences = text.split('. ')
            if summary_type == 'brief':
                return '. '.join(sentences[:3]) + '.'
            elif summary_type == 'bullet_points':
                return "\n• " + "\n• ".join(sentences[:5])
            else:
                return '. '.join(sentences[:7]) + '.'

        model = get_summarizer()

        # Adjust length based on type
        if summary_type == 'brief':
            max_length = 100
            min_length = 30
        elif summary_type == 'detailed':
            max_length = 300
            min_length = 100
        else:  # bullet_points
            max_length = 150
            min_length = 50

        # Handle long texts by chunking
        max_chunk_size = 1024
        if len(text) > max_chunk_size:
            sentences = text.split('. ')
            chunks = []
            current_chunk = ""

            for sentence in sentences:
                if len(current_chunk) + len(sentence) < max_chunk_size:
                    current_chunk += " " + sentence
                else:
                    if current_chunk:
                        chunks.append(current_chunk)
                    current_chunk = sentence

            if current_chunk:
                chunks.append(current_chunk)

            summaries = []
            for chunk in chunks[:5]:
                summary = model(chunk, max_length=max_length, min_length=min_length, do_sample=False)
                summaries.append(summary[0]['summary_text'])

            final_summary = " ".join(summaries)
        else:
            summary = model(text, max_length=max_length, min_length=min_length, do_sample=False)
            final_summary = summary[0]['summary_text']

        # Format as bullet points if requested
        if summary_type == 'bullet_points':
            sentences = final_summary.split('. ')
            final_summary = "\n• " + "\n• ".join(sentences)

        return final_summary
    except Exception as e:
        print(f"Error generating summary: {e}")
        # Fallback to simple summary
        sentences = text.split('. ')
        if summary_type == 'brief':
            return '. '.join(sentences[:3])
        elif summary_type == 'bullet_points':
            return "\n• " + "\n• ".join(sentences[:5])
        else:
            return '. '.join(sentences[:7])

def extract_action_items(text):
    """Extract potential action items from text"""
    action_verbs = ['need to', 'should', 'must', 'have to', 'will', 'going to', 'plan to', 'todo', 'action', 'follow up']
    
    sentences = sent_tokenize(text)
    action_items = []
    
    for sentence in sentences:
        sentence_lower = sentence.lower()
        if any(verb in sentence_lower for verb in action_verbs):
            action_items.append(sentence.strip())
    
    return action_items[:10]  # Limit to 10 items

def analyze_sentiment(text):
    """Analyze sentiment of the text"""
    try:
        analyzer = get_sentiment_analyzer()
        
        # Analyze in chunks if text is too long
        max_length = 512
        if len(text) > max_length:
            chunks = [text[i:i+max_length] for i in range(0, len(text), max_length)]
            results = []
            for chunk in chunks[:5]:  # Limit to 5 chunks
                result = analyzer(chunk)[0]
                results.append(result)
            
            # Average sentiment
            positive_count = sum(1 for r in results if r['label'] == 'POSITIVE')
            avg_score = sum(r['score'] for r in results) / len(results)
            
            return {
                'label': 'POSITIVE' if positive_count > len(results) / 2 else 'NEGATIVE',
                'score': avg_score,
                'details': results
            }
        else:
            result = analyzer(text)[0]
            return {
                'label': result['label'],
                'score': result['score'],
                'details': [result]
            }
    except Exception as e:
        print(f"Error in sentiment analysis: {e}")
        return {'label': 'NEUTRAL', 'score': 0.5, 'details': []}

def translate_text(text, target_language):
    """Translate text to target language - supports deep-translator and googletrans"""
    if not GOOGLETRANS_AVAILABLE:
        return None
    try:
        # Use deep-translator (recommended - no httpx conflicts)
        if DEEP_TRANSLATOR_AVAILABLE:
            translated = GoogleTranslator(source='auto', target=target_language).translate(text)
            return {
                'text': translated,
                'source_language': 'auto',
                'target_language': target_language,
                'confidence': 0.9
            }
        # Fallback to googletrans
        elif translator is not None:
            translation = translator.translate(text, dest=target_language)
            return {
                'text': translation.text,
                'source_language': translation.src,
                'target_language': target_language,
                'confidence': 0.9
            }
        return None
    except Exception as e:
        print(f"Error translating text: {e}")
        return None

def create_docx_export(data, output_path):
    """Create DOCX export with transcript and summary"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    doc = Document()
    
    # Title
    title = doc.add_heading('Voice Summary Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_heading('Document Information', 2)
    metadata_table = doc.add_table(rows=4, cols=2)
    metadata_table.style = 'Light Grid Accent 1'
    
    cells = metadata_table.rows[0].cells
    cells[0].text = 'Filename'
    cells[1].text = data.get('filename', 'Unknown')
    
    cells = metadata_table.rows[1].cells
    cells[0].text = 'Duration'
    cells[1].text = f"{data.get('duration', 0):.2f} seconds"
    
    cells = metadata_table.rows[2].cells
    cells[0].text = 'Language'
    cells[1].text = data.get('language', 'Unknown')
    
    cells = metadata_table.rows[3].cells
    cells[0].text = 'Processed Date'
    cells[1].text = data.get('timestamp', 'Unknown')
    
    doc.add_paragraph()
    
    # Summary
    doc.add_heading('Summary', 2)
    doc.add_paragraph(data.get('summary', 'No summary available'))
    
    # Keywords
    if data.get('keywords'):
        doc.add_heading('Key Terms', 2)
        keywords_text = ', '.join([kw['word'] for kw in data['keywords'][:10]])
        doc.add_paragraph(keywords_text)
    
    # Action Items
    if data.get('action_items'):
        doc.add_heading('Action Items', 2)
        for item in data['action_items']:
            doc.add_paragraph(item, style='List Bullet')
    
    # Full Transcript
    doc.add_page_break()
    doc.add_heading('Full Transcript', 2)
    doc.add_paragraph(data.get('transcript', 'No transcript available'))
    
    doc.save(output_path)

def create_pdf_export(data, output_path):
    """Create PDF export with transcript and summary"""
    if not REPORTLAB_AVAILABLE:
        raise ImportError("reportlab not installed. Run: pip install reportlab")
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#2C3E50'),
        spaceAfter=30,
        alignment=1  # Center
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#34495E'),
        spaceAfter=12,
        spaceBefore=12
    )
    
    # Title
    story.append(Paragraph("Voice Summary Report", title_style))
    story.append(Spacer(1, 0.3*inch))
    
    # Metadata table
    metadata = [
        ['Filename', data.get('filename', 'Unknown')],
        ['Duration', f"{data.get('duration', 0):.2f} seconds"],
        ['Language', data.get('language', 'Unknown')],
        ['Processed Date', data.get('timestamp', 'Unknown')]
    ]
    
    t = Table(metadata, colWidths=[2*inch, 4*inch])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#ECF0F1')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#BDC3C7'))
    ]))
    story.append(t)
    story.append(Spacer(1, 0.3*inch))
    
    # Summary
    story.append(Paragraph("Summary", heading_style))
    story.append(Paragraph(data.get('summary', 'No summary available'), styles['BodyText']))
    story.append(Spacer(1, 0.2*inch))
    
    # Keywords
    if data.get('keywords'):
        story.append(Paragraph("Key Terms", heading_style))
        keywords_text = ', '.join([kw['word'] for kw in data['keywords'][:10]])
        story.append(Paragraph(keywords_text, styles['BodyText']))
        story.append(Spacer(1, 0.2*inch))
    
    # Action Items
    if data.get('action_items'):
        story.append(Paragraph("Action Items", heading_style))
        for item in data['action_items']:
            story.append(Paragraph(f"• {item}", styles['BodyText']))
        story.append(Spacer(1, 0.2*inch))
    
    # Full Transcript
    story.append(PageBreak())
    story.append(Paragraph("Full Transcript", heading_style))
    story.append(Paragraph(data.get('transcript', 'No transcript available'), styles['BodyText']))
    
    doc.build(story)

def create_srt_export(segments, output_path):
    """Create SRT subtitle file from transcript segments"""
    with open(output_path, 'w', encoding='utf-8') as f:
        for i, segment in enumerate(segments, 1):
            start_time = format_srt_time(segment['start'])
            end_time = format_srt_time(segment['end'])
            text = segment['text'].strip()
            
            f.write(f"{i}\n")
            f.write(f"{start_time} --> {end_time}\n")
            f.write(f"{text}\n\n")

def format_srt_time(seconds):
    """Format seconds to SRT time format (HH:MM:SS,mmm)"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    millis = int((seconds % 1) * 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"

def save_to_history(data):
    """Save processing history"""
    try:
        if os.path.exists(HISTORY_FILE):
            with open(HISTORY_FILE, 'r') as f:
                history = json.load(f)
        else:
            history = []
        
        history.append(data)
        
        # Keep only last 100 entries
        history = history[-100:]
        
        with open(HISTORY_FILE, 'w') as f:
            json.dump(history, f, indent=2)
    except Exception as e:
        print(f"Error saving to history: {e}")

# API Routes

@app.route('/')
def serve_index():
    """Serve the frontend HTML directly from Flask"""
    return app.send_static_file('index.html')

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({'status': 'ok', 'message': 'Server is running'})

@app.route('/api/transcribe', methods=['POST'])
def transcribe_audio():
    """Main transcription endpoint with all features"""

    print(f"\n=== NEW TRANSCRIPTION REQUEST ===")
    print(f"Files in request: {list(request.files.keys())}")
    print(f"Form data: {dict(request.form)}")

    if 'file' not in request.files:
        print("ERROR: No file in request")
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']
    print(f"File received: {file.filename}, size: {file.content_length}")

    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    if not allowed_file(file.filename):
        return jsonify({'error': f'Invalid file format. Allowed: mp3, wav, ogg, flac, m4a'}), 400

    try:
        print(f"Starting processing for: {file.filename}")
        # Get options from request
        options = {
            'summary_type': request.form.get('summary_type', 'brief'),
            'include_keywords': request.form.get('include_keywords', 'true') == 'true',
            'include_sentiment': request.form.get('include_sentiment', 'true') == 'true',
            'include_action_items': request.form.get('include_action_items', 'true') == 'true',
            'translate_to': request.form.get('translate_to', None)
        }
        
        # Save uploaded file
        file_id = str(uuid.uuid4())
        filename = f"{file_id}_{file.filename}"
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        file.save(filepath)

        # Emit progress update
        socketio.emit('progress', {'stage': 'audio_analysis', 'progress': 10})

        # Extract audio features
        audio_features = extract_audio_features(filepath)

        # ✅ Validate audio before sending to Whisper
        duration = audio_features.get('duration', 0)
        if duration < 0.5:
            return jsonify({
                'error': 'Audio file is too short or silent. Please upload an audio file with at least 1 second of speech.'
            }), 400

        # ✅ Check audio is not silent using librosa
        if LIBROSA_AVAILABLE and np is not None:
            try:
                y, sr = librosa.load(filepath, sr=16000)
                # Check if audio has any meaningful content
                if len(y) == 0:
                    return jsonify({'error': 'Audio file appears to be empty.'}), 400
                rms = float(np.sqrt(np.mean(y ** 2)))
                if rms < 0.0001:
                    return jsonify({'error': 'Audio file is silent. Please upload a file with actual speech.'}), 400
                # Pad audio if it is very short to avoid reshape error
                min_samples = sr * 1  # at least 1 second
                if len(y) < min_samples:
                    y = np.pad(y, (0, min_samples - len(y)))
                    # Save padded version back
                    import soundfile as sf
                    sf.write(filepath, y, sr)
            except Exception as e:
                print(f"Audio validation warning: {e}")

        socketio.emit('progress', {'stage': 'transcription', 'progress': 30})

        # ✅ Transcribe with Whisper - with error handling
        try:
            model = get_whisper_model()
            result = model.transcribe(
                filepath,
                task='transcribe',
                fp16=False,          # Disable half precision - avoids reshape issues
                language=None,       # Auto detect language
                verbose=False
            )
        except RuntimeError as e:
            error_msg = str(e)
            if 'reshape' in error_msg or 'shape' in error_msg or 'elements' in error_msg:
                return jsonify({
                    'error': 'Audio file is too short or silent. Please upload a longer audio file with clear speech (at least 2-3 seconds).'
                }), 400
            raise e

        transcript = result['text'].strip()
        segments = result.get('segments', [])
        detected_language = result.get('language', 'unknown')

        # ✅ Check transcript is not empty
        if not transcript:
            return jsonify({
                'error': 'No speech detected in the audio file. Please check your file has clear spoken audio.'
            }), 400
        
        socketio.emit('progress', {'stage': 'language_detection', 'progress': 50})
        
        # Advanced language detection
        language_details = detect_language_advanced(transcript)
        
        socketio.emit('progress', {'stage': 'summarization', 'progress': 60})
        
        # Generate summary
        summary = generate_summary(transcript, options['summary_type'])
        
        # Extract keywords
        keywords = extract_keywords(transcript) if options['include_keywords'] else []
        
        socketio.emit('progress', {'stage': 'analysis', 'progress': 75})
        
        # Extract action items
        action_items = extract_action_items(transcript) if options['include_action_items'] else []
        
        # Sentiment analysis
        sentiment = analyze_sentiment(transcript) if options['include_sentiment'] else None
        
        socketio.emit('progress', {'stage': 'translation', 'progress': 85})
        
        # Translation if requested
        translation = None
        if options['translate_to']:
            translation = translate_text(summary, options['translate_to'])
        
        socketio.emit('progress', {'stage': 'finalizing', 'progress': 95})
        
        # Prepare response data
        timestamp = datetime.now().isoformat()
        
        response_data = {
            'id': file_id,
            'filename': file.filename,
            'timestamp': timestamp,
            'transcript': transcript,
            'summary': summary,
            'language': detected_language,
            'language_details': language_details,
            'duration': audio_features['duration'],
            'audio_features': audio_features,
            'segments': segments,
            'keywords': keywords,
            'action_items': action_items,
            'sentiment': sentiment,
            'translation': translation,
            'options': options
        }
        
        # Save to history
        save_to_history({
            'id': file_id,
            'filename': file.filename,
            'timestamp': timestamp,
            'language': detected_language,
            'duration': audio_features['duration']
        })
        
        socketio.emit('progress', {'stage': 'complete', 'progress': 100})
        
        return jsonify(response_data)
        
    except Exception as e:
        print(f"Error processing file: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/export/<file_id>', methods=['POST'])
def export_file(file_id):
    """Export transcript in various formats"""
    
    try:
        data = request.json
        export_format = data.get('format', 'txt')
        export_data = data.get('data', {})
        
        output_filename = f"{file_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{export_format}"
        output_path = os.path.join(OUTPUT_FOLDER, output_filename)
        
        if export_format == 'txt':
            # Plain text export
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(f"Transcript: {export_data.get('filename', 'Unknown')}\n")
                f.write(f"Date: {export_data.get('timestamp', 'Unknown')}\n")
                f.write(f"Duration: {export_data.get('duration', 0):.2f}s\n")
                f.write(f"Language: {export_data.get('language', 'Unknown')}\n\n")
                f.write("=" * 50 + "\n")
                f.write("SUMMARY\n")
                f.write("=" * 50 + "\n\n")
                f.write(export_data.get('summary', '') + "\n\n")
                f.write("=" * 50 + "\n")
                f.write("FULL TRANSCRIPT\n")
                f.write("=" * 50 + "\n\n")
                f.write(export_data.get('transcript', ''))
        
        elif export_format == 'docx':
            create_docx_export(export_data, output_path)
        
        elif export_format == 'pdf':
            create_pdf_export(export_data, output_path)
        
        elif export_format == 'srt':
            segments = export_data.get('segments', [])
            create_srt_export(segments, output_path)
        
        elif export_format == 'json':
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False)
        
        else:
            return jsonify({'error': 'Unsupported export format'}), 400
        
        return send_file(output_path, as_attachment=True, download_name=output_filename)
        
    except Exception as e:
        print(f"Error exporting file: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/history', methods=['GET'])
def get_history():
    """Get processing history"""
    try:
        if os.path.exists(HISTORY_FILE):
            with open(HISTORY_FILE, 'r') as f:
                history = json.load(f)
            return jsonify(history)
        else:
            return jsonify([])
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/translate', methods=['POST'])
def translate():
    """Translate text endpoint"""
    try:
        data = request.json
        text = data.get('text', '')
        target_language = data.get('target_language', 'en')
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        translation = translate_text(text, target_language)
        
        if translation:
            return jsonify(translation)
        else:
            return jsonify({'error': 'Translation failed'}), 500
            
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/languages', methods=['GET'])
def get_supported_languages():
    """Get list of supported languages"""
    languages = {
        'en': 'English',
        'es': 'Spanish',
        'fr': 'French',
        'de': 'German',
        'it': 'Italian',
        'pt': 'Portuguese',
        'ru': 'Russian',
        'ja': 'Japanese',
        'ko': 'Korean',
        'zh-cn': 'Chinese (Simplified)',
        'zh-tw': 'Chinese (Traditional)',
        'ar': 'Arabic',
        'hi': 'Hindi',
        'nl': 'Dutch',
        'pl': 'Polish',
        'tr': 'Turkish',
        'sv': 'Swedish',
        'da': 'Danish',
        'no': 'Norwegian',
        'fi': 'Finnish'
    }
    return jsonify(languages)

# WebSocket events for real-time updates
@socketio.on('connect')
def handle_connect():
    print('Client connected')
    emit('connection_response', {'status': 'connected'})

@socketio.on('disconnect')
def handle_disconnect():
    print('Client disconnected')

if __name__ == '__main__':
    print("\n" + "="*50)
    print("  Voice Summarizer Pro")
    print("="*50)
    print(f"\n  Open in browser: http://localhost:5000")
    print(f"  API endpoint:    http://localhost:5000/api")
    print(f"\n  DO NOT open index.html directly!")
    print(f"  Use the URL above instead.")
    print("\n" + "="*50 + "\n")
    socketio.run(app, debug=True, host='0.0.0.0', port=5000, use_reloader=True, log_output=True)